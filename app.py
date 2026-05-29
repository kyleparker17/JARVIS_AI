import streamlit as st
from datetime import datetime
from docx import Document
import io
import os
import pandas as pd
from pypdf import PdfReader

from langchain_ollama import ChatOllama

st.set_page_config(page_title="JARVIS", layout="wide", initial_sidebar_state="expanded")

st.markdown("""
<style>
    .stApp {
        background: radial-gradient(1100px 600px at 18% -10%, #12141d 0%, #0a0a0f 55%);
        color: #e6e8ef;
    }
    h1 {
        background: linear-gradient(90deg, #00ff9d 0%, #00d4ff 100%);
        -webkit-background-clip: text; -webkit-text-fill-color: transparent;
        font-weight: 800; letter-spacing: -1.5px; margin-bottom: 0;
    }
    .chat-message {
        padding: 13px 19px; border-radius: 20px; margin-bottom: 10px;
        max-width: 78%; line-height: 1.65; font-size: 0.95rem;
        box-shadow: 0 3px 14px rgba(0,0,0,0.35);
        animation: fadeIn 0.22s ease;
    }
    .user-message {
        background: linear-gradient(135deg, #1f2433 0%, #181c28 100%);
        border: 1px solid rgba(0,212,255,0.18);
        border-bottom-right-radius: 5px; margin-left: auto;
    }
    .assistant-message {
        background: #1c1e27; border: 1px solid rgba(0,255,157,0.15);
        border-bottom-left-radius: 5px;
    }
    @keyframes fadeIn {
        from { opacity:0; transform:translateY(5px); }
        to   { opacity:1; transform:translateY(0); }
    }
    .stTextInput input {
        background-color: #15171f !important; border: 1px solid #2b2e3a !important;
        border-radius: 14px !important; color: #e6e8ef !important; padding: 11px 15px !important;
    }
    .stTextInput input:focus {
        border-color: #00ff9d !important;
        box-shadow: 0 0 0 2px rgba(0,255,157,0.15) !important;
    }
    div[data-baseweb="select"] > div {
        background-color: #15171f !important; border: 1px solid #2b2e3a !important;
        border-radius: 14px !important;
    }
    .stButton > button { border-radius: 14px; border: 1px solid #2b2e3a; font-weight: 600; transition: all 0.15s ease; }
    button[data-testid="stBaseButton-primary"] {
        background: linear-gradient(135deg, #00ff9d 0%, #00d4ff 100%);
        color: #07120d; border: none;
    }
    [data-testid="stSidebar"] { background-color: #0d0f16; border-right: 1px solid #1d2030; }
    /* 전체 여백 축소 */
    .block-container { padding: 1rem 2rem 0.5rem !important; max-width: 100% !important; }
    section[data-testid="stSidebar"] > div { padding: 0.6rem 0.7rem !important; }
    h1 { margin-top: 0 !important; margin-bottom: 0 !important; }
    /* 사이드바 버튼 슬림 */
    [data-testid="stSidebar"] .stButton > button {
        font-size: 0.8rem !important; padding: 3px 10px !important;
        font-weight: 400 !important; border-radius: 8px !important;
        min-height: 0 !important; line-height: 1.6 !important;
    }
    [data-testid="stSidebar"] hr { margin: 5px 0 !important; }
    [data-testid="stSidebar"] p,
    [data-testid="stSidebar"] small { font-size: 0.8rem !important; margin: 0 !important; }
    /* 스크롤바 */
    ::-webkit-scrollbar { width: 6px; }
    ::-webkit-scrollbar-thumb { background: #2b2e3a; border-radius: 6px; }
    ::-webkit-scrollbar-track { background: transparent; }
</style>
""", unsafe_allow_html=True)

st.title("JARVIS")
st.caption("v7 • Multi-Agent System")

# ====================== Multi-Chat + Watch Folder (사이드바 전용 추가) ======================
if "watch_path" not in st.session_state:
    st.session_state.watch_path = ""
if "watch_files" not in st.session_state:
    st.session_state.watch_files = []
if "watch_mtimes" not in st.session_state:
    st.session_state.watch_mtimes = {}

if "chats" not in st.session_state:
    cid = "c_" + str(int(datetime.now().timestamp()))
    st.session_state.chats = {
        cid: {
            "id": cid,
            "title": "New Chat",
            "messages": [],
            "attached_files": [],
            "watched_folder": None,
            "folder_files": []
        }
    }
    st.session_state.current_chat_id = cid

def get_current_chat():
    return st.session_state.chats[st.session_state.current_chat_id]

def create_new_chat():
    cid = "c_" + str(int(datetime.now().timestamp()))
    st.session_state.chats[cid] = {
        "id": cid,
        "title": "New Chat",
        "messages": [],
        "attached_files": [],
        "watched_folder": None,
        "folder_files": []
    }
    st.session_state.current_chat_id = cid
    st.session_state.chat_input = ""
    return cid

def _scan_folder(path):
    """폴더 스캔 + mtime 기록. 반환: (files list, mtimes dict)"""
    exts = (".pdf", ".docx", ".txt", ".md", ".csv", ".xlsx")
    files, mtimes = [], {}
    for root, _, fnames in os.walk(path):
        for fname in fnames:
            if not fname.lower().endswith(exts):
                continue
            full = os.path.join(root, fname)
            try:
                mtimes[full] = os.path.getmtime(full)
                ext = fname.lower()
                if ext.endswith(".pdf"):
                    with open(full, "rb") as f:
                        content = "\n".join(p.extract_text() or "" for p in PdfReader(f).pages)
                elif ext.endswith(".docx"):
                    content = "\n".join(p.text for p in Document(full).paragraphs)
                else:
                    with open(full, "r", encoding="utf-8", errors="ignore") as f:
                        content = f.read()
                files.append({"name": fname, "path": full, "content": content[:4000]})
            except:
                pass
    return files, mtimes

def extract_text_from_file(uploaded_file):
    name = uploaded_file.name.lower()
    try:
        if name.endswith(".pdf"):
            reader = PdfReader(uploaded_file)
            return "\n".join([page.extract_text() or "" for page in reader.pages])
        elif name.endswith(".docx"):
            doc = Document(uploaded_file)
            return "\n".join([p.text for p in doc.paragraphs])
        elif name.endswith((".txt", ".md")):
            return uploaded_file.getvalue().decode("utf-8", errors="ignore")
        elif name.endswith(".csv"):
            return pd.read_csv(uploaded_file).to_string()
        elif name.endswith((".xlsx", ".xls")):
            return pd.read_excel(uploaded_file).to_string()
        else:
            return uploaded_file.getvalue().decode("utf-8", errors="ignore")
    except Exception as e:
        return f"[Error reading file: {str(e)}]"

# ====================== 사이드바 (여기에만 Multi-Agent + 채팅목록 + Watch Folder) ======================
with st.sidebar:
    st.header("🧠 Multi-Agent System")
    
    current = get_current_chat()
    model = st.session_state.get("active_orchestrator", "mistral-small:24b")
    if model == "mistral-small:24b":
        st.markdown('<span style="background:#00ff9d; color:#000; padding:5px 14px; border-radius:9999px; font-size:0.85em; font-weight:700;">MAIN ORCHESTRATOR</span> **Mistral Small 24B**', unsafe_allow_html=True)
    else:
        st.markdown('<span style="background:#ffaa00; color:#000; padding:5px 14px; border-radius:9999px; font-size:0.85em; font-weight:700;">2ND ORCHESTRATOR</span> **Phi-4 14B**', unsafe_allow_html=True)
    
    st.divider()
    st.write("**Agent Pool**")
    st.caption("• llama3.1:8b → 문서 / RAG / Watch Folder")
    st.caption("• gemma2:9b → 코드 실행 / 데이터 분석")
    
    st.divider()
    st.session_state.multi_agent_mode = st.toggle("Multi-Agent Mode", value=st.session_state.get("multi_agent_mode", True))

    if st.button("🔄 Orchestrator 전환 (24B ↔ 14B)"):
        st.session_state.active_orchestrator = "phi4:14b" if model == "mistral-small:24b" else "mistral-small:24b"
        st.rerun()

    # === 채팅목록 + 새채팅 만들기 ===
    st.divider()
    st.subheader("💬 Chats")
    if st.button("＋ New Chat", use_container_width=True):
        create_new_chat()
        st.rerun()

    for cid, chat in sorted(st.session_state.chats.items(), key=lambda x: x[1].get("created_at", ""), reverse=True):
        is_active = cid == st.session_state.current_chat_id
        title = chat["title"][:20] + "..." if len(chat["title"]) > 20 else chat["title"]
        label = ("● " + title) if is_active else title
        if st.button(label, key=f"chat_{cid}", use_container_width=True):
            st.session_state.current_chat_id = cid
            st.session_state.chat_input = ""
            st.rerun()

    # === Watch Folder (전역 고정) ===
    st.divider()
    st.caption("📁 Watch Folder")
    wf_input = st.text_input("", placeholder="폴더 경로 붙여넣기",
                              value=st.session_state.watch_path,
                              label_visibility="collapsed", key="wf_input")
    wc1, wc2 = st.columns(2)
    with wc1:
        if st.button("연결", use_container_width=True):
            if wf_input and os.path.exists(wf_input):
                st.session_state.watch_path = wf_input
                st.session_state.watch_files, st.session_state.watch_mtimes = _scan_folder(wf_input)
                st.rerun()
            else:
                st.error("경로 없음")
    with wc2:
        if st.button("재스캔", use_container_width=True, disabled=not st.session_state.watch_path):
            st.session_state.watch_files, st.session_state.watch_mtimes = _scan_folder(st.session_state.watch_path)
            st.rerun()
    if st.session_state.watch_path:
        st.caption(f"✅ {len(st.session_state.watch_files)}개 · {st.session_state.watch_path[-28:]}")
        for f in st.session_state.watch_files[:6]:
            st.caption(f"· {f['name']}")
        if len(st.session_state.watch_files) > 6:
            st.caption(f"  +{len(st.session_state.watch_files)-6}개")

if "composer_nonce" not in st.session_state:
    st.session_state.composer_nonce = 0

SKILLS = {
    "기본":     "You are JARVIS, a helpful multi-agent IT assistant.",
    "데이터 분석": "You are a senior data analyst. Report clear, actionable insights.",
    "문서 요약":  "You summarize documents into concise, well-structured Korean summaries.",
    "코드 작성":  "You are a senior Python engineer. Write safe, production-ready code.",
    "전략 보고서": "You are an IT strategy consultant producing high-quality 2026-trend reports.",
}

# ====================== 채팅 영역 (스크롤 고정 박스) ======================
current_chat = get_current_chat()

chat_box = st.container(height=560, border=True)
with chat_box:
    for msg in current_chat["messages"]:
        if msg["role"] == "user":
            with st.chat_message("user", avatar="👤"):
                st.markdown(f'<div class="chat-message user-message">{msg["content"]}</div>', unsafe_allow_html=True)
        else:
            with st.chat_message("assistant", avatar="🔎"):
                st.markdown(f'<div class="chat-message assistant-message">{msg["content"]}</div>', unsafe_allow_html=True)

# ====================== Composer (하단) — 파일+스킬 → 채팅창 → 모델 → 전송 ======================
col_attach, col_input, col_model, col_send = st.columns(
    [0.08, 0.64, 0.20, 0.08], vertical_alignment="center"
)

with col_attach:
    with st.popover("➕", use_container_width=True):
        st.caption("파일 + 스킬")
        uploaded_file = st.file_uploader(
            "파일 업로드", type=["pdf","docx","txt","md","csv","xlsx"],
            key="composer_file"
        )
        selected_skill = st.selectbox("스킬", list(SKILLS.keys()), key="composer_skill")

with col_input:
    prompt = st.text_input(
        "메시지", placeholder="Message JARVIS...", label_visibility="collapsed",
        key=f"chat_input_{st.session_state.composer_nonce}"
    )

with col_model:
    active_model = st.selectbox(
        "모델", ["mistral-small:24b", "phi4:14b"], index=0,
        label_visibility="collapsed", key="composer_model"
    )

with col_send:
    send = st.button("전송", use_container_width=True, type="primary")

# ====================== 입력 처리 ======================
if send and prompt and prompt.strip():
    # mtime 변경 감지 → 자동 재스캔 (실시간급)
    if st.session_state.watch_path and st.session_state.watch_mtimes:
        if any(os.path.getmtime(p) != st.session_state.watch_mtimes.get(p, 0)
               for p in list(st.session_state.watch_mtimes)):
            st.session_state.watch_files, st.session_state.watch_mtimes = _scan_folder(st.session_state.watch_path)

    ctx = ""
    if uploaded_file:
        ctx += f"\n\n[Attached: {uploaded_file.name}]\n{extract_text_from_file(uploaded_file)[:4000]}"
    if st.session_state.watch_files:
        ctx += "\n\n[Watch Folder]\n" + "\n".join(
            [f"--- {f['name']} ---\n{f['content']}" for f in st.session_state.watch_files]
        )

    current_chat["messages"].append({"role": "user", "content": prompt})

    try:
        llm = ChatOllama(model=active_model, temperature=0.25, num_predict=8192)
        system = SKILLS.get(selected_skill, SKILLS["기본"])
        if st.session_state.get("multi_agent_mode", True):
            system += " Coordinate your expert agents to give the best possible answer."
        full_response = llm.invoke(system + ctx + "\n\nUser: " + prompt).content
    except Exception as e:
        full_response = f"오류: {e}"

    current_chat["messages"].append({"role": "assistant", "content": full_response})
    if current_chat["title"] == "New Chat" and len(current_chat["messages"]) == 2:
        current_chat["title"] = prompt[:28]
    st.session_state.composer_nonce += 1
    st.rerun()
